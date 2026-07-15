"""문서 파일에서 텍스트를 읽어온다. (.pdf / .docx / .txt / .md)"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md", ".markdown"}

# 폴더 안내문이 본문 문서로 인덱싱되어 검색 결과를 오염시키는 것을 막는다.
IGNORED_NAMES = {"readme.md", "readme.txt", "readme"}


@dataclass
class LoadedSection:
    """파일에서 읽어낸 텍스트 한 덩어리. page는 PDF에서만 채워진다."""

    source: str
    text: str
    page: int | None = None


class UnsupportedFileError(ValueError):
    pass


def load_path(path: Path) -> list[LoadedSection]:
    """파일 하나를 읽어 섹션 목록으로 돌려준다."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        sections = _load_pdf(path)
    elif suffix == ".docx":
        sections = _load_docx(path)
    elif suffix in {".txt", ".md", ".markdown"}:
        sections = _load_text(path)
    else:
        raise UnsupportedFileError(
            f"지원하지 않는 형식입니다: {suffix or '(확장자 없음)'} "
            f"— 지원 형식: {', '.join(sorted(SUPPORTED_SUFFIXES))}"
        )
    return [s for s in sections if s.text.strip()]


def load_directory(directory: Path) -> list[LoadedSection]:
    """폴더를 재귀적으로 훑어 지원 형식 파일을 모두 읽는다."""
    sections: list[LoadedSection] = []
    for path in iter_supported_files(directory):
        sections.extend(load_path(path))
    return sections


def iter_supported_files(directory: Path) -> Iterable[Path]:
    """인덱싱 대상 파일 목록.

    Word 임시 파일(~$...)과 폴더 안내용 README, 숨김 파일은 제외한다.
    """
    if not directory.exists():
        return []
    return sorted(p for p in directory.rglob("*") if _is_indexable(p))


def _is_indexable(path: Path) -> bool:
    if not path.is_file() or path.suffix.lower() not in SUPPORTED_SUFFIXES:
        return False
    if path.name.lower() in IGNORED_NAMES:
        return False
    # ~$취업규칙.docx — Word가 파일을 열어둔 동안 만드는 잠금 파일.
    return not path.name.startswith((".", "~$"))


def _load_pdf(path: Path) -> list[LoadedSection]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return [
        LoadedSection(source=path.name, text=_normalize(page.extract_text() or ""), page=i)
        for i, page in enumerate(reader.pages, start=1)
    ]


def _load_docx(path: Path) -> list[LoadedSection]:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]

    # 표 안의 내용도 규정·매뉴얼류에서는 본문 못지않게 중요한 경우가 많다.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return [LoadedSection(source=path.name, text=_normalize("\n".join(parts)))]


def _load_text(path: Path) -> list[LoadedSection]:
    for encoding in ("utf-8", "utf-8-sig", "cp949"):
        try:
            raw = path.read_text(encoding=encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        raw = path.read_text(encoding="utf-8", errors="replace")
    return [LoadedSection(source=path.name, text=_normalize(raw))]


def _normalize(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
